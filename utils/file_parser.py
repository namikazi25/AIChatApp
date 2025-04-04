import os
import io
from fastapi import UploadFile
from PIL import Image
from pypdf import PdfReader
import docx

async def parse_file(file: UploadFile) -> dict:
    """Parse different file types and extract their content."""
    content = await file.read()
    _, ext = os.path.splitext(file.filename)
    ext = ext.lower()

    try:
        if ext in {".jpg", ".jpeg", ".png"}:
            return await parse_image(content)
        elif ext == ".pdf":
            text = await parse_pdf(content)
            return {"type": "text", "content": text}
        elif ext in {".docx", ".doc"}:
            text = await parse_docx(content)
            return {"type": "text", "content": text}
        else:
            return {"type": "error", "content": f"Unsupported file type: {ext}. Please upload JPG, PNG, PDF, or DOCX files."}
    except Exception as e:
        return {"type": "error", "content": f"Error parsing file: {str(e)}"}
    finally:
        await file.seek(0)

async def parse_image(content: bytes) -> dict:
    """Process image file and return both metadata and base64 encoded image."""
    try:
        import base64
        with Image.open(io.BytesIO(content)) as image:
            width, height = image.size
            format_name = image.format
            mode = image.mode
            
            # Convert image to base64 for multimodal models
            buffered = io.BytesIO()
            image.save(buffered, format=format_name)
            img_str = base64.b64encode(buffered.getvalue()).decode('utf-8')
            
            # Return both metadata and base64 encoded image
            return {
                "type": "image",
                "metadata": f"Image: {width}x{height} pixels, {format_name} format, {mode} mode.",
                "base64": img_str,
                "mime_type": f"image/{format_name.lower()}"
            }
    except Exception as e:
        return {
            "type": "error",
            "content": f"Error processing image: {str(e)}"
        }

async def parse_pdf(content: bytes) -> str:
    """Extract text from a PDF file using pypdf."""
    try:
        pdf = PdfReader(io.BytesIO(content))
        text = ""
        for page in pdf.pages:
            text += page.extract_text() or ""
        return text.strip() if text.strip() else "This PDF appears to be a scanned document with no extractable text."
    except Exception as e:
        return f"Error processing PDF: {str(e)}"

async def parse_docx(content: bytes) -> str:
    """Extract text from a DOCX file."""
    try:
        doc = docx.Document(io.BytesIO(content))
        text = "\n".join(para.text for para in doc.paragraphs)

        for table in doc.tables:
            for row in table.rows:
                text += "\n" + " ".join(cell.text for cell in row.cells)

        return text.strip()
    except Exception as e:
        return f"Error processing DOCX: {str(e)}"
